"""
app/resume/renderer.py — PDF/DOCX/Markdown/JSON export from ResumeData.

DESIGN PRINCIPLE: The renderer consumes ResumeData JSON.
The LLM NEVER generates formatting or layout.
All templates are deterministic Python code.

Templates:
  - modern    (default — colorful, clean, ATS-safe)
  - classic_ats (pure ATS-optimized, no color)
  - minimal   (ultra-clean whitespace)
  - executive (dark accent, C-suite tone)
  - developer (monospace accents, GitHub-style)
  - academic  (traditional, citations)
"""

import html
import io
from typing import Literal, Optional, List

from app.resume.models import ResumeData, ExperienceEntry, EducationEntry

TemplateType = Literal["classic_ats", "modern", "minimal", "executive", "developer", "academic"]


def _esc(text: Optional[str]) -> str:
    """Safely escape text for ReportLab Paragraph XML parser."""
    if not text:
        return ""
    return html.escape(text, quote=False)


# ─────────────────────────────────────────────────────────────────────────────
#  PDF Generation (reportlab)
# ─────────────────────────────────────────────────────────────────────────────

def _get_template_colors(template: TemplateType) -> dict:
    colors = {
        "modern":      {"primary": (0.18, 0.25, 0.75), "accent": (0.20, 0.62, 0.86), "text": (0.1, 0.1, 0.1), "light": (0.94, 0.95, 1.0)},
        "classic_ats": {"primary": (0.1, 0.1, 0.1),  "accent": (0.3, 0.3, 0.3),  "text": (0.0, 0.0, 0.0), "light": (0.95, 0.95, 0.95)},
        "minimal":     {"primary": (0.2, 0.2, 0.2),  "accent": (0.5, 0.5, 0.5),  "text": (0.1, 0.1, 0.1), "light": (0.97, 0.97, 0.97)},
        "executive":   {"primary": (0.05, 0.15, 0.3), "accent": (0.7, 0.55, 0.2), "text": (0.05, 0.05, 0.05), "light": (0.93, 0.95, 0.97)},
        "developer":   {"primary": (0.05, 0.5, 0.35), "accent": (0.1, 0.8, 0.5),  "text": (0.05, 0.1, 0.05), "light": (0.93, 0.98, 0.95)},
        "academic":    {"primary": (0.3, 0.1, 0.5),   "accent": (0.5, 0.2, 0.7),  "text": (0.05, 0.0, 0.1),  "light": (0.96, 0.94, 0.98)},
    }
    return colors.get(template, colors["modern"])


class PDFDependencyError(Exception):
    """Raised when reportlab dependency is missing on the server for PDF export."""
    pass


def generate_pdf(resume: ResumeData, template: TemplateType = "modern") -> bytes:
    """Generate a beautiful PDF from ResumeData. Returns bytes."""
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
            Table, TableStyle, KeepTogether
        )
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    except ImportError:
        raise PDFDependencyError(
            "PDF generation library (reportlab) is not installed on the server. Please install reportlab or export in JSON, LaTeX, or Markdown format."
        )

    buffer = io.BytesIO()
    tc = _get_template_colors(template)
    primary = colors.Color(*tc["primary"])
    accent = colors.Color(*tc["accent"])
    text_color = colors.Color(*tc["text"])
    light_bg = colors.Color(*tc["light"])

    doc = SimpleDocTemplate(
        buffer,
        pagesize=LETTER,
        rightMargin=0.6 * inch,
        leftMargin=0.6 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
    )

    story = []
    styles = getSampleStyleSheet()

    # ── Custom Styles ─────────────────────────────────────────────────────────
    name_style = ParagraphStyle(
        "Name", fontSize=22, fontName="Helvetica-Bold",
        textColor=primary, spaceAfter=2, leading=26,
    )
    headline_style = ParagraphStyle(
        "Headline", fontSize=11, fontName="Helvetica",
        textColor=accent, spaceAfter=4, leading=14,
    )
    contact_style = ParagraphStyle(
        "Contact", fontSize=8.5, fontName="Helvetica",
        textColor=text_color, spaceAfter=0, leading=12,
    )
    section_style = ParagraphStyle(
        "Section", fontSize=10.5, fontName="Helvetica-Bold",
        textColor=primary, spaceBefore=8, spaceAfter=2,
        leading=13, borderPadding=(0, 0, 2, 0),
    )
    body_style = ParagraphStyle(
        "Body", fontSize=9, fontName="Helvetica",
        textColor=text_color, spaceAfter=2, leading=13,
    )
    bullet_style = ParagraphStyle(
        "Bullet", fontSize=9, fontName="Helvetica",
        textColor=text_color, spaceAfter=1, leading=13,
        leftIndent=12, bulletIndent=0,
    )
    subhead_style = ParagraphStyle(
        "Subhead", fontSize=9.5, fontName="Helvetica-Bold",
        textColor=text_color, spaceAfter=0, leading=12,
    )
    sub2_style = ParagraphStyle(
        "Sub2", fontSize=8.5, fontName="Helvetica-Oblique",
        textColor=colors.Color(0.4, 0.4, 0.4), spaceAfter=2, leading=11,
    )
    skill_tag_style = ParagraphStyle(
        "SkillTag", fontSize=8.5, fontName="Helvetica",
        textColor=text_color, leading=11,
    )

    def section_header(title: str):
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=1.2, color=primary, spaceAfter=4))

    # ── Header ────────────────────────────────────────────────────────────────
    p = resume.personal
    story.append(Paragraph(p.name or "Your Name", name_style))

    if resume.headline:
        story.append(Paragraph(resume.headline, headline_style))

    # Contact info row
    contact_parts = []
    if p.email:
        contact_parts.append(p.email)
    if p.phone:
        contact_parts.append(p.phone)
    if p.location:
        contact_parts.append(p.location)
    if p.linkedin:
        contact_parts.append(p.linkedin.replace("https://", ""))
    if p.github:
        contact_parts.append(p.github.replace("https://", ""))
    if p.website:
        contact_parts.append(p.website.replace("https://", ""))

    if contact_parts:
        contact_text = "  •  ".join(_esc(cp) for cp in contact_parts)
        story.append(Paragraph(contact_text, contact_style))

    story.append(Spacer(1, 6))

    # ── Summary ───────────────────────────────────────────────────────────────
    if resume.summary:
        section_header("Professional Summary")
        story.append(Paragraph(_esc(resume.summary), body_style))

    # ── Skills ───────────────────────────────────────────────────────────────
    if resume.skills:
        section_header("Skills")
        skill_rows = []
        for sg in resume.skills:
            if not sg.skills:
                continue
            cat_para = Paragraph(f"<b>{_esc(sg.category)}</b>", skill_tag_style)
            skills_para = Paragraph(", ".join(_esc(s) for s in sg.skills), skill_tag_style)
            skill_rows.append([cat_para, skills_para])

        if skill_rows:
            t = Table(skill_rows, colWidths=[1.3 * inch, 5.6 * inch])
            t.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 1),
            ]))
            story.append(t)

    # ── Experience ────────────────────────────────────────────────────────────
    if resume.experience:
        section_header("Experience")
        for exp in resume.experience:
            date_str = f"{_esc(exp.start_date)} – {_esc(exp.end_date) or 'Present'}"
            # Role + Company row with date on right
            role_para = Paragraph(f"<b>{_esc(exp.role)}</b>", subhead_style)
            date_para = Paragraph(date_str, ParagraphStyle(
                "DateRight", fontSize=8.5, fontName="Helvetica",
                textColor=colors.Color(0.4, 0.4, 0.4), alignment=TA_RIGHT, leading=12,
            ))
            row = Table(
                [[role_para, date_para]],
                colWidths=[4.5 * inch, 2.4 * inch]
            )
            row.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
            ]))
            story.append(row)
            story.append(Paragraph(f"{_esc(exp.company)}" + (f" — {_esc(exp.location)}" if exp.location else ""), sub2_style))
            for bullet in exp.bullets:
                story.append(Paragraph(f"• {_esc(bullet)}", bullet_style))
            story.append(Spacer(1, 4))

    # ── Projects ─────────────────────────────────────────────────────────────
    if resume.projects:
        section_header("Projects")
        for proj in resume.projects:
            tech_str = ", ".join(_esc(t) for t in proj.technologies[:8]) if proj.technologies else ""
            title_para = Paragraph(
                f"<b>{_esc(proj.name)}</b>" + (f" | <i>{tech_str}</i>" if tech_str else ""),
                subhead_style
            )
            story.append(title_para)
            if proj.description:
                story.append(Paragraph(_esc(proj.description), body_style))
            for bullet in proj.bullets:
                story.append(Paragraph(f"• {_esc(bullet)}", bullet_style))
            story.append(Spacer(1, 4))

    # ── Education ─────────────────────────────────────────────────────────────
    if resume.education:
        section_header("Education")
        for edu in resume.education:
            date_str = f"{edu.start_date} – {edu.end_date}" if edu.start_date else edu.end_date
            deg_str = f"{edu.degree}" + (f" in {edu.field_of_study}" if edu.field_of_study else "")
            inst_para = Paragraph(f"<b>{edu.institution}</b>", subhead_style)
            date_para = Paragraph(date_str or "", ParagraphStyle(
                "DateRight2", fontSize=8.5, fontName="Helvetica",
                textColor=colors.Color(0.4, 0.4, 0.4), alignment=TA_RIGHT, leading=12,
            ))
            row = Table([[inst_para, date_para]], colWidths=[4.5 * inch, 2.4 * inch])
            row.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 0),
                ("TOPPADDING", (0, 0), (-1, -1), 0),
            ]))
            story.append(row)
            story.append(Paragraph(deg_str + (f" | GPA: {edu.gpa}" if edu.gpa else ""), sub2_style))
            story.append(Spacer(1, 4))

    # ── Certifications ────────────────────────────────────────────────────────
    if resume.certifications:
        section_header("Certifications")
        for cert in resume.certifications:
            cert_str = f"<b>{cert.name}</b>" + (f" — {cert.issuer}" if cert.issuer else "") + (f" ({cert.date})" if cert.date else "")
            story.append(Paragraph(cert_str, body_style))

    # ── Languages ─────────────────────────────────────────────────────────────
    if resume.languages:
        section_header("Languages")
        lang_str = "  •  ".join(
            f"{l.language} ({l.proficiency})" if l.proficiency else l.language
            for l in resume.languages
        )
        story.append(Paragraph(lang_str, body_style))

    doc.build(story)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX Generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(resume: ResumeData) -> bytes:
    """Generate a Word DOCX from ResumeData."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise RuntimeError("python-docx is required. Install: pip install python-docx")

    doc = Document()
    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    p = resume.personal

    # Name
    h = doc.add_heading(p.name or "Your Name", 0)
    h.runs[0].font.color.rgb = RGBColor(0x2E, 0x40, 0xBF)

    # Contact
    contact_parts = []
    if p.email:
        contact_parts.append(p.email)
    if p.phone:
        contact_parts.append(p.phone)
    if p.location:
        contact_parts.append(p.location)
    if p.linkedin:
        contact_parts.append(p.linkedin)
    if p.github:
        contact_parts.append(p.github)

    if contact_parts:
        c = doc.add_paragraph(" | ".join(contact_parts))
        c.runs[0].font.size = Pt(9)

    if resume.headline:
        hl = doc.add_paragraph(resume.headline)
        hl.runs[0].font.italic = True

    def add_section(title: str):
        h = doc.add_heading(title, 2)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x2E, 0x40, 0xBF)

    if resume.summary:
        add_section("Professional Summary")
        doc.add_paragraph(resume.summary)

    if resume.skills:
        add_section("Skills")
        for sg in resume.skills:
            p_para = doc.add_paragraph()
            r = p_para.add_run(f"{sg.category}: ")
            r.font.bold = True
            p_para.add_run(", ".join(sg.skills))

    if resume.experience:
        add_section("Experience")
        for exp in resume.experience:
            role_para = doc.add_paragraph()
            r = role_para.add_run(exp.role)
            r.font.bold = True
            role_para.add_run(f" — {exp.company}" + (f", {exp.location}" if exp.location else ""))
            date_para = doc.add_paragraph(f"{exp.start_date} – {exp.end_date or 'Present'}")
            date_para.runs[0].font.italic = True
            date_para.runs[0].font.size = Pt(9)
            for bullet in exp.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if resume.projects:
        add_section("Projects")
        for proj in resume.projects:
            p_para = doc.add_paragraph()
            r = p_para.add_run(proj.name)
            r.font.bold = True
            if proj.technologies:
                p_para.add_run(f" | {', '.join(proj.technologies[:6])}")
            if proj.description:
                doc.add_paragraph(proj.description)
            for bullet in proj.bullets:
                doc.add_paragraph(bullet, style="List Bullet")

    if resume.education:
        add_section("Education")
        for edu in resume.education:
            e_para = doc.add_paragraph()
            r = e_para.add_run(edu.institution)
            r.font.bold = True
            deg = f"{edu.degree}" + (f" in {edu.field_of_study}" if edu.field_of_study else "")
            doc.add_paragraph(deg + (f" | GPA: {edu.gpa}" if edu.gpa else ""))

    if resume.certifications:
        add_section("Certifications")
        for cert in resume.certifications:
            doc.add_paragraph(
                cert.name + (f" — {cert.issuer}" if cert.issuer else "") + (f" ({cert.date})" if cert.date else "")
            )

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
#  Markdown Export
# ─────────────────────────────────────────────────────────────────────────────

def generate_markdown(resume: ResumeData) -> str:
    """Generate Markdown from ResumeData."""
    lines = []
    p = resume.personal

    lines.append(f"# {p.name or 'Resume'}")
    if resume.headline:
        lines.append(f"\n_{resume.headline}_")

    contact = " | ".join(filter(None, [p.email, p.phone, p.location, p.linkedin, p.github, p.website]))
    if contact:
        lines.append(f"\n{contact}")

    if resume.summary:
        lines.append("\n## Summary\n")
        lines.append(resume.summary)

    if resume.skills:
        lines.append("\n## Skills\n")
        for sg in resume.skills:
            lines.append(f"**{sg.category}:** {', '.join(sg.skills)}")

    if resume.experience:
        lines.append("\n## Experience\n")
        for exp in resume.experience:
            lines.append(f"### {exp.role} — {exp.company}")
            lines.append(f"_{exp.start_date} – {exp.end_date or 'Present'}_ | {exp.location}")
            for b in exp.bullets:
                lines.append(f"- {b}")
            lines.append("")

    if resume.projects:
        lines.append("\n## Projects\n")
        for proj in resume.projects:
            lines.append(f"### {proj.name}")
            if proj.technologies:
                lines.append(f"**Tech:** {', '.join(proj.technologies)}")
            if proj.description:
                lines.append(proj.description)
            for b in proj.bullets:
                lines.append(f"- {b}")
            lines.append("")

    if resume.education:
        lines.append("\n## Education\n")
        for edu in resume.education:
            lines.append(f"### {edu.institution}")
            deg = f"{edu.degree}" + (f" in {edu.field_of_study}" if edu.field_of_study else "")
            lines.append(deg + (f" | GPA: {edu.gpa}" if edu.gpa else ""))
            if edu.start_date or edu.end_date:
                start = edu.start_date or ''
                end = edu.end_date or ''
                lines.append(f"_{start} – {end}_")
            lines.append("")

    if resume.certifications:
        lines.append("\n## Certifications\n")
        for cert in resume.certifications:
            lines.append(f"- **{cert.name}**" + (f" — {cert.issuer}" if cert.issuer else ""))

    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
#  LaTeX Code Generation
# ─────────────────────────────────────────────────────────────────────────────

import re

def escape_latex(text: str) -> str:
    """
    Escape special LaTeX characters in standard user text fields.
    Order is crucial: backslash must be escaped first!
    """
    if not text:
        return ""
    
    # Strip 4-byte unicode emojis that break standard pdfLaTeX engines
    text = re.sub(r'[\U00010000-\U0010FFFF]', '', text)

    replacements = [
        ('\\', r'\textbackslash{}'),
        ('&', r'\&'),
        ('%', r'\%'),
        ('$', r'\$'),
        ('#', r'\#'),
        ('_', r'\_'),
        ('{', r'\{'),
        ('}', r'\}'),
        ('~', r'\textasciitilde{}'),
        ('^', r'\textasciicircum{}'),
        ('<', r'\textless{}'),
        ('>', r'\textgreater{}'),
    ]
    res = str(text)
    for old, new in replacements:
        res = res.replace(old, new)

    # Convert multiline formatting from textareas
    res = res.replace('\r\n', '\n')
    res = res.replace('\n\n', r' \par ')
    res = res.replace('\n', r' \\ ')
    return res


def escape_latex_url(url: str) -> str:
    """
    Escape URLs for hyperref \\href{url}{text}.
    Only escape % (to \\%) since it is a LaTeX comment delimiter.
    Leave & as-is: it is a valid URL query-parameter separator and
    hyperref handles it correctly inside \\href{} argument.
    Ensures absolute URL protocol (https://) is present for web targets.
    """
    if not url:
        return ""
    res = str(url).strip()
    if not (res.startswith("http://") or res.startswith("https://") or res.startswith("mailto:")):
        res = "https://" + res
    res = res.replace('\\', '%5C')
    res = res.replace('%', r'\%')
    # Do NOT escape & — hyperref accepts raw & inside \href{} URL argument
    return res




def generate_latex(resume: ResumeData, template: TemplateType = "modern") -> str:
    """
    Generate compilation-ready, structured LaTeX code from ResumeData.
    Supports 6 template variations with customized palettes & typography.
    """
    p = resume.personal

    # Template palette config
    palettes = {
        "modern":      {"primary": "2E40C7", "accent": "339EDC", "text": "1A1A1A"},
        "classic_ats": {"primary": "1F1F1F", "accent": "4A4A4A", "text": "000000"},
        "minimal":     {"primary": "333333", "accent": "666666", "text": "1A1A1A"},
        "executive":   {"primary": "0D264D", "accent": "B28C33", "text": "0D0D0D"},
        "developer":   {"primary": "0D8059", "accent": "1ACD80", "text": "0D1A0D"},
        "academic":    {"primary": "4C197F", "accent": "8033B2", "text": "0D001A"},
    }
    pal = palettes.get(template, palettes["modern"])

    lines = []
    lines.append(r"\documentclass[10pt,a4paper]{article}")
    lines.append(r"\usepackage[utf8]{inputenc}")
    lines.append(r"\usepackage[margin=0.6in]{geometry}")
    lines.append(r"\usepackage{hyperref}")
    lines.append(r"\usepackage{titlesec}")
    lines.append(r"\usepackage{enumitem}")
    lines.append(r"\usepackage{xcolor}")
    lines.append(r"\usepackage{fontawesome5}")
    
    if template == "developer":
        lines.append(r"\usepackage{inconsolata}")

    # Color definitions
    lines.append(f"\\definecolor{{primary}}{{HTML}}{{{pal['primary']}}}")
    lines.append(f"\\definecolor{{accent}}{{HTML}}{{{pal['accent']}}}")
    lines.append(f"\\definecolor{{darktext}}{{HTML}}{{{pal['text']}}}")

    # Hyperref setup
    lines.append(r"\hypersetup{colorlinks=true, linkcolor=primary, urlcolor=primary}")

    # Page setup
    lines.append(r"\pagestyle{empty}")
    lines.append(r"\setlength{\parindent}{0pt}")
    lines.append(r"\setlist[itemize]{leftmargin=*, noitemsep, topsep=2pt, parsep=1pt}")

    # Section formatting
    lines.append(r"\titleformat{\section}{\large\bfseries\color{primary}\uppercase}{}{0em}{}[\titlerule]")
    lines.append(r"\titlespacing*{\section}{0pt}{10pt}{4pt}")

    lines.append(r"\begin{document}")

    # ── Header Section ────────────────────────────────────────────────────────
    name_esc = escape_latex(p.name or "Your Name")
    lines.append(r"\begin{center}")
    lines.append(f"{{\\Huge \\bfseries \\color{{primary}} {name_esc}}}")

    if resume.headline:
        head_esc = escape_latex(resume.headline)
        lines.append(f"\\\\[4pt]{{\\large \\color{{accent}} \\textit{{{head_esc}}}}}")

    # Contact items row
    contacts = []
    if p.email:
        email_clean = escape_latex_url(p.email)
        email_txt = escape_latex(p.email)
        contacts.append(f"\\faEnvelope\\ \\href{{mailto:{email_clean}}}{{{email_txt}}}")
    if p.phone:
        phone_txt = escape_latex(p.phone)
        contacts.append(f"\\faPhone\\ {phone_txt}")
    if p.location:
        loc_txt = escape_latex(p.location)
        contacts.append(f"\\faMapMarkerAlt\\ {loc_txt}")
    if p.linkedin:
        link_clean = escape_latex_url(p.linkedin)
        link_txt = escape_latex(p.linkedin.replace("https://", ""))
        contacts.append(f"\\faLinkedin\\ \\href{{{link_clean}}}{{{link_txt}}}")
    if p.github:
        gh_clean = escape_latex_url(p.github)
        gh_txt = escape_latex(p.github.replace("https://", ""))
        contacts.append(f"\\faGithub\\ \\href{{{gh_clean}}}{{{gh_txt}}}")
    if p.website:
        web_clean = escape_latex_url(p.website)
        web_txt = escape_latex(p.website.replace("https://", ""))
        contacts.append(f"\\faGlobe\\ \\href{{{web_clean}}}{{{web_txt}}}")

    if contacts:
        contact_str = " \\ \\$\\cdot\\$ \\ ".join(contacts)
        lines.append(f"\\\\[6pt]{{\\small \\color{{darktext}} {contact_str}}}")

    lines.append(r"\end{center}")
    lines.append(r"\vspace{-6pt}")

    # ── Summary Section ───────────────────────────────────────────────────────
    if resume.summary:
        lines.append(r"\section*{Professional Summary}")
        lines.append(f"{escape_latex(resume.summary)}")

    # ── Skills Section ────────────────────────────────────────────────────────
    if resume.skills:
        lines.append(r"\section*{Skills}")
        lines.append(r"\begin{itemize}[label={}]")
        for sg in resume.skills:
            if sg.skills:
                cat = escape_latex(sg.category)
                sk_list = escape_latex(", ".join(sg.skills))
                lines.append(f"  \\item \\textbf{{{cat}:}} {sk_list}")
        lines.append(r"\end{itemize}")

    # ── Work Experience Section ───────────────────────────────────────────────
    if resume.experience:
        lines.append(r"\section*{Work Experience}")
        for exp in resume.experience:
            role_esc = escape_latex(exp.role)
            comp_esc = escape_latex(exp.company)
            loc_esc = escape_latex(exp.location)
            date_str = escape_latex(f"{exp.start_date} -- {exp.end_date or 'Present'}")

            lines.append(r"\noindent")
            lines.append(f"\\textbf{{{role_esc}}} \\hfill {{\\small \\color{{darktext}} {date_str}}}")
            lines.append(r"\\")
            lines.append(f"\\textit{{{comp_esc}}}" + (f", {loc_esc}" if loc_esc else ""))

            if exp.bullets:
                lines.append(r"\begin{itemize}")
                for b in exp.bullets:
                    lines.append(f"  \\item {escape_latex(b)}")
                lines.append(r"\end{itemize}")
            lines.append(r"\vspace{4pt}")

    # ── Projects Section ──────────────────────────────────────────────────────
    if resume.projects:
        lines.append(r"\section*{Projects}")
        for proj in resume.projects:
            name_esc = escape_latex(proj.name)
            tech_esc = escape_latex(", ".join(proj.technologies)) if proj.technologies else ""

            proj_link = getattr(proj, "link", None) or proj.url or proj.github_url
            link_latex = ""
            if proj_link:
                url_target = escape_latex_url(proj_link)
                display_txt = escape_latex(proj_link.replace("https://", "").replace("http://", ""))
                link_latex = f" \\ \\href{{{url_target}}}{{\\small \\color{{accent}} \\faLink\\ {display_txt}}}"

            lines.append(r"\noindent")
            lines.append(f"\\textbf{{{name_esc}}}" + link_latex + (f" \\hfill {{\\small \\color{{accent}} ({tech_esc})}}" if tech_esc else ""))

            if proj.description:
                lines.append(f"\\\\ {escape_latex(proj.description)}")

            if proj.bullets:
                lines.append(r"\begin{itemize}")
                for b in proj.bullets:
                    lines.append(f"  \\item {escape_latex(b)}")
                lines.append(r"\end{itemize}")
            lines.append(r"\vspace{4pt}")


    # ── Education Section ─────────────────────────────────────────────────────
    if resume.education:
        lines.append(r"\section*{Education}")
        for edu in resume.education:
            inst_esc = escape_latex(edu.institution)
            deg_esc = escape_latex(edu.degree)
            field_esc = escape_latex(edu.field_of_study) if edu.field_of_study else ""
            date_str = escape_latex(f"{edu.start_date} -- {edu.end_date}")
            gpa_str = escape_latex(f"GPA: {edu.gpa}") if edu.gpa else ""

            lines.append(r"\noindent")
            lines.append(f"\\textbf{{{inst_esc}}} \\hfill {{\\small \\color{{darktext}} {date_str}}}")
            lines.append(r"\\")
            deg_line = f"\\textit{{{deg_esc}}}" + (f" in {field_esc}" if field_esc else "") + (f" ({gpa_str})" if gpa_str else "")
            lines.append(deg_line)
            lines.append(r"\vspace{4pt}")

    # ── Certifications Section ────────────────────────────────────────────────
    if resume.certifications:
        lines.append(r"\section*{Certifications}")
        lines.append(r"\begin{itemize}")
        for cert in resume.certifications:
            c_name = escape_latex(cert.name)
            c_issuer = escape_latex(cert.issuer) if cert.issuer else ""
            c_date = escape_latex(cert.date) if cert.date else ""
            item_str = f"\\textbf{{{c_name}}}" + (f" -- {c_issuer}" if c_issuer else "") + (f" ({c_date})" if c_date else "")
            lines.append(f"  \\item {item_str}")
        lines.append(r"\end{itemize}")

    # ── Achievements Section ──────────────────────────────────────────────────
    if resume.achievements:
        lines.append(r"\section*{Achievements}")
        lines.append(r"\begin{itemize}")
        for ach in resume.achievements:
            t = (ach.title or "").strip().lstrip("–—•- ").strip()
            d = (ach.description or "").strip().lstrip("–—•- ").strip()

            t_esc = escape_latex(t)
            d_esc = escape_latex(d)

            if d and d != t:
                if t:
                    lines.append(f"  \\item \\textbf{{{t_esc}}}: {d_esc}")
                else:
                    lines.append(f"  \\item {d_esc}")
            elif t:
                lines.append(f"  \\item \\textbf{{{t_esc}}}")
        lines.append(r"\end{itemize}")


    lines.append(r"\end{document}")
    return "\n".join(lines)


# ─────────────────────────────────────────────────────────────────────────────
#  Public Export API
# ─────────────────────────────────────────────────────────────────────────────

def export_resume(
    resume: ResumeData,
    format: str = "pdf",
    template: TemplateType = "modern",
) -> tuple[bytes, str, str]:
    """
    Export resume in the requested format.
    Returns (content_bytes, media_type, filename_extension).
    """
    name_slug = (resume.personal.name or "resume").lower().replace(" ", "_")[:20]

    if format == "pdf":
        content = generate_pdf(resume, template)
        return content, "application/pdf", f"{name_slug}_resume.pdf"

    elif format == "docx":
        content = generate_docx(resume)
        return content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{name_slug}_resume.docx"

    elif format == "markdown":
        md = generate_markdown(resume)
        return md.encode("utf-8"), "text/markdown", f"{name_slug}_resume.md"

    elif format == "json":
        json_str = resume.model_dump_json(indent=2)
        return json_str.encode("utf-8"), "application/json", f"{name_slug}_resume.json"

    elif format in ("latex", "tex"):
        tex_code = generate_latex(resume, template)
        return tex_code.encode("utf-8"), "text/x-tex", f"{name_slug}_resume.tex"

    raise ValueError(f"Unsupported export format: {format}")

